from typing import Dict, Optional

from bs4.element import Comment, NavigableString, Tag
from docx.oxml.ns import qn
from docx.shared import Pt

from .html_docx_constants import _BLOCK_TAGS, _INHERITABLE_STYLES, _INLINE_TAGS
from .html_docx_css import _parse_background_color, _parse_border_left, _parse_padding
from .html_docx_images import _add_image_run, _add_svg_run
from .html_docx_paragraphs import (
    _apply_paragraph_style,
    _apply_run_style,
    _apply_ul_margin_indent,
    _resolve_list_indent_pt,
)
from .html_docx_selectors import _compute_style_map


_SKIP_TAGS = {"style", "script", "head", "meta", "link", "title", "noscript"}


def _handle_block(node, target, css_rules, parent_style: Dict[str, str], table_auto_widths) -> None:
    target_container = _ensure_container(target)
    if isinstance(node, Comment):
        return
    if isinstance(node, Tag) and node.name in _SKIP_TAGS:
        return
    if isinstance(node, NavigableString):
        text = node.strip()
        if not text:
            return
        paragraph = _ensure_paragraph(target_container)
        _apply_paragraph_style(paragraph, parent_style)
        run = paragraph.add_run(_transform_text(_normalize_text(text), parent_style))
        _apply_run_style(run, parent_style)
        return

    if not isinstance(node, Tag):
        return

    if node.name == "img":
        paragraph = _ensure_paragraph(target_container)
        _apply_paragraph_style(paragraph, parent_style)
        _add_image_run(paragraph, node, parent_style)
        return

    if node.name == "svg":
        paragraph = _ensure_paragraph(target_container)
        _apply_paragraph_style(paragraph, parent_style)
        _add_svg_run(paragraph, node, parent_style)
        return

    if node.name == "table":
        from .html_docx_tables import _handle_table

        _handle_table(node, target_container, css_rules, parent_style, table_auto_widths)
        return

    if node.name in _BLOCK_TAGS:
        current_style = _merge_styles(parent_style, _compute_style_map(node, css_rules))

        if _should_wrap_container(node, current_style):
            container = _add_container(
                target_container, node, current_style, css_rules, table_auto_widths
            )
            for child in node.children:
                _handle_block(child, container, css_rules, current_style, table_auto_widths)
            return

        if _has_block_children(node):
            for child in node.children:
                _handle_block(child, target_container, css_rules, current_style, table_auto_widths)
            return

        paragraph = _ensure_paragraph(
            target_container,
            style="List Bullet" if node.name == "li" else None,
        )
        _apply_paragraph_style(paragraph, current_style)
        if node.name == "li":
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            # Apply the parent <ul> margin-left as additional left indent so the
            # list is offset the same amount it is in the HTML rendering.
            _apply_ul_margin_indent(paragraph, _resolve_list_indent_pt(parent_style))
        _add_inline_runs(node, paragraph, css_rules, current_style)
        return

    if node.name in _INLINE_TAGS:
        paragraph = _ensure_paragraph(target_container)
        _add_inline_runs(node, paragraph, css_rules, parent_style)
        return

    for child in node.children:
        _handle_block(child, target_container, css_rules, parent_style, table_auto_widths)


def _add_inline_runs(node, paragraph, css_rules, parent_style: Dict[str, str]) -> None:
    if isinstance(node, Comment):
        return
    if isinstance(node, NavigableString):
        raw_text = str(node)
        if not raw_text.strip():
            paragraph.add_run(" ")
            return
        normalized = _normalize_inline_text(raw_text)
        if not paragraph.text and normalized.startswith(" "):
            normalized = normalized.lstrip()
        if not normalized:
            return
        run = paragraph.add_run(_transform_text(normalized, parent_style))
        _apply_run_style(run, parent_style)
        return

    if not isinstance(node, Tag):
        return

    if node.name == "img":
        _add_image_run(paragraph, node, parent_style)
        return

    if node.name == "svg":
        _add_svg_run(paragraph, node, parent_style)
        return

    if node.name == "br":
        paragraph.add_run().add_break()
        return

    current_style = _merge_styles(parent_style, _compute_style_map(node, css_rules))

    if node.name in _INLINE_TAGS:
        if node.name in {"strong", "b"}:
            current_style = {**current_style, "font-weight": "bold"}
        if node.name in {"em", "i"}:
            current_style = {**current_style, "font-style": "italic"}
        if node.name == "u":
            current_style = {**current_style, "text-decoration": "underline"}

    for child in node.children:
        _add_inline_runs(child, paragraph, css_rules, current_style)


def _transform_text(text: str, style_map: Dict[str, str]) -> str:
    transform = style_map.get("text-transform", "").strip().lower()
    if transform == "uppercase":
        return text.upper()
    if transform == "lowercase":
        return text.lower()
    if transform == "capitalize":
        return " ".join(word.capitalize() for word in text.split())
    return text


def _has_block_children(node: Tag) -> bool:
    for child in node.children:
        if isinstance(child, Tag) and child.name in _BLOCK_TAGS:
            return True
    return False


def _merge_styles(parent_style: Dict[str, str], own_style: Dict[str, str]) -> Dict[str, str]:
    merged: Dict[str, str] = {}
    for key in _INHERITABLE_STYLES:
        if key in parent_style:
            merged[key] = parent_style[key]
    for key, value in own_style.items():
        merged[key] = value
    return merged


def _should_wrap_container(node: Tag, style_map: Dict[str, str]) -> bool:
    if style_map.get("display", "").strip().lower() == "flex":
        return True
    if _parse_background_color(style_map) is not None:
        return True
    if _parse_border_left(style_map) is not None:
        return True
    padding = _parse_padding(style_map.get("padding", ""))
    if padding and any(v > 0 for v in padding):
        return True
    if node.name in {"header", "section"}:
        return True
    return False


def _add_container(target, node: Tag, style_map: Dict[str, str], css_rules, table_auto_widths):
    from .html_docx_tables import _apply_cell_styles, _apply_table_styles

    if style_map.get("display", "").strip().lower() == "flex":
        children = [child for child in node.children if isinstance(child, Tag)]
        if len(children) == 2:
            table = target.add_table(rows=1, cols=2)
            _apply_table_styles(table, style_map)
            left_cell = table.cell(0, 0)
            right_cell = table.cell(0, 1)
            _apply_cell_styles(left_cell, style_map)
            _apply_cell_styles(right_cell, style_map)
            for child in children[0].children:
                _handle_block(child, left_cell, css_rules, style_map, table_auto_widths)
            for child in children[1].children:
                _handle_block(child, right_cell, css_rules, style_map, table_auto_widths)
            return table

    table = target.add_table(rows=1, cols=1)
    _apply_table_styles(table, style_map)
    cell = table.cell(0, 0)
    _apply_cell_styles(cell, style_map)
    return cell


def _ensure_container(target):
    if hasattr(target, "add_paragraph"):
        return target
    if hasattr(target, "cell"):
        return target.cell(0, 0)
    return target


def _ensure_paragraph(container, style: Optional[str] = None):
    if hasattr(container, "paragraphs") and container.paragraphs:
        last = container.paragraphs[-1]
        if not last.text and len(last.runs) == 0 and _is_pristine_paragraph(last):
            if style is not None:
                last.style = style
            return last
    return container.add_paragraph(style=style)


def _is_pristine_paragraph(paragraph) -> bool:
    """Return True only for the unformatted initial placeholder paragraph.

    We must not reuse paragraphs that carry visual formatting (borders, exact
    spacing, shading) from a previously processed block element, otherwise the
    next element would overwrite that formatting.  A pristine paragraph is one
    whose pPr contains only a pStyle declaration (or nothing at all).
    """
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return True
    _ALLOWED_TAGS = {qn("w:pStyle"), qn("w:rPr")}
    return all(child.tag in _ALLOWED_TAGS for child in p_pr)


def _normalize_text(text: str) -> str:
    return " ".join(text.split())


def _normalize_inline_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    if not text:
        return ""
    lead_space = text[:1].isspace()
    trail_space = text[-1:].isspace()
    collapsed = " ".join(text.split())
    if not collapsed and (lead_space or trail_space):
        return " "
    if lead_space:
        collapsed = " " + collapsed
    if trail_space:
        collapsed = collapsed + " "
    return collapsed
