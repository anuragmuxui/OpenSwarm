from typing import Dict, List, Optional, Tuple

from bs4.element import Tag
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .html_docx_constants import _INHERITABLE_STYLES, _NAMED_COLORS, _PADDING_SCALE
from .html_docx_css import (
    _border_sz,
    _normalize_padding,
    _parse_background_color,
    _parse_border,
    _parse_border_left,
    _parse_percentage,
    _parse_px_to_pt,
    _resolve_padding,
)
from .html_docx_paragraphs import (
    _add_list_indent_padding,
    _add_paragraph_indent,
    _add_paragraph_spacing,
)
from .html_docx_playwright import _extract_auto_widths
from .html_docx_selectors import _compute_style_map
from .html_docx_shared import _remove_trailing_empty_paragraph


def _handle_table(
    table_node: Tag, target_container, css_rules, parent_style: Dict[str, str], table_auto_widths
) -> None:
    table_style = _merge_styles(parent_style, _compute_style_map(table_node, css_rules))
    parent_padding = _resolve_padding(parent_style)
    parent_padding_right = None
    if parent_padding:
        _top, right, _bottom, _left = _normalize_padding(parent_padding)
        parent_padding_right = right if right else None
    rows = table_node.find_all("tr", recursive=False)
    if not rows:
        for section in table_node.find_all(["thead", "tbody", "tfoot"], recursive=False):
            rows.extend(section.find_all("tr", recursive=False))
    if not rows:
        return

    max_cols = 0
    row_cells: list[list[Tag]] = []
    row_colspans: list[list[int]] = []
    for row in rows:
        cells = [cell for cell in row.find_all(["td", "th"], recursive=False)]
        if not cells:
            continue
        colspans = []
        col_count = 0
        for cell in cells:
            colspan = cell.get("colspan")
            try:
                span_value = int(colspan) if colspan else 1
            except ValueError:
                span_value = 1
            span_value = max(span_value, 1)
            colspans.append(span_value)
            col_count += span_value
        max_cols = max(max_cols, col_count)
        row_cells.append(cells)
        row_colspans.append(colspans)

    if max_cols == 0:
        return

    _remove_trailing_empty_paragraph(target_container)
    docx_table = target_container.add_table(rows=len(rows), cols=max_cols)
    _apply_table_styles(docx_table, table_style)
    _apply_table_parent_padding(docx_table, parent_style)
    if _should_prevent_row_split(table_node, table_style):
        _set_table_cant_split(docx_table)
    column_widths_pt = _extract_table_column_widths(table_node, table_style, max_cols, table_auto_widths)
    if not column_widths_pt:
        column_widths_pt = _extract_auto_widths(table_node, table_auto_widths, max_cols)
    if column_widths_pt:
        _apply_table_column_widths(docx_table, column_widths_pt)

    is_collapsed = table_style.get("border-collapse", "").strip().lower() == "collapse"
    collapsed_borders: dict = {}
    if is_collapsed:
        collapsed_borders = _collect_collapsed_borders(
            row_cells, row_colspans, max_cols, css_rules
        )

    total_rows = len(row_cells)
    for row_idx, cells in enumerate(row_cells):
        col_idx = 0
        for cell_idx, cell_node in enumerate(cells):
            if col_idx >= max_cols:
                break
            docx_cell = docx_table.cell(row_idx, col_idx)
            colspan = (
                row_colspans[row_idx][cell_idx] if cell_idx < len(row_colspans[row_idx]) else 1
            )
            if colspan > 1 and col_idx + colspan - 1 < max_cols:
                docx_cell = docx_cell.merge(docx_table.cell(row_idx, col_idx + colspan - 1))
                if column_widths_pt and colspan == max_cols and len(cells) == 1:
                    span_width = sum(column_widths_pt)
                    if span_width > 0:
                        _set_cell_width(docx_cell, span_width)

            cell_own_style = _compute_style_map(cell_node, css_rules)
            cell_style = _merge_styles(table_style, cell_own_style)
            # CSS `background` is not inherited, but it shows through transparent cells
            # visually. In DOCX every cell is opaque by default (white), so we must
            # explicitly forward the background to any cell that doesn't declare its own.
            # We also fall back to parent_style so that cells inside a nested table that
            # sits inside a coloured cell also receive the correct background.
            effective_bg = _parse_background_color(table_style) or _parse_background_color(
                parent_style
            )
            if effective_bg and not _parse_background_color(cell_own_style):
                # Store with '#' prefix so _parse_background_color can re-read it
                # when this cell_style is passed as parent_style to nested elements.
                cell_style = {**cell_style, "background-color": f"#{effective_bg}"}
            if cell_node.name == "th":
                cell_style = {**cell_style, "font-weight": "bold"}
            if (
                parent_padding_right is not None
                and cell_style.get("text-align", "").strip().lower() == "right"
                and not cell_style.get("padding-right")
                and not cell_style.get("padding")
            ):
                cell_style = {**cell_style, "padding-right": f"{parent_padding_right}pt"}

            # For border-collapse:collapse tables, interior borders are handled by
            # tblBorders/insideH+insideV (applied once after the loop). Suppress cell-level
            # borders on interior edges so they don't double up with the table-level lines.
            # Outer-edge cells keep their cell-level borders (one side each).
            suppress_borders: Optional[set] = None
            if is_collapsed:
                is_first_row = row_idx == 0
                is_last_row = row_idx == total_rows - 1
                is_first_col = col_idx == 0
                is_last_col = col_idx + max(colspan, 1) - 1 >= max_cols - 1
                suppress_borders = set()
                if not is_first_row:
                    suppress_borders.add("top")
                if not is_last_row:
                    suppress_borders.add("bottom")
                if not is_first_col:
                    suppress_borders.add("left")
                if not is_last_col:
                    suppress_borders.add("right")

            _apply_cell_styles(docx_cell, cell_style, suppress_borders=suppress_borders)
            if (
                len(cells) == 2
                and cell_style.get("text-align", "").strip().lower() == "right"
            ):
                _set_cell_no_wrap(docx_cell)
                _set_cell_width(docx_cell, _estimate_right_column_width_pt(column_widths_pt))
            if column_widths_pt and col_idx < len(column_widths_pt):
                col_width_pt = column_widths_pt[col_idx]
                _set_cell_width(docx_cell, col_width_pt)
                # Inject absolute cell width so nested image sizing can cap to it.
                cell_style = {**cell_style, "_cell_width_pt": str(col_width_pt)}

            from .html_docx_blocks import (
                _add_inline_runs,
                _handle_block,
                _has_block_children,
            )
            from .html_docx_paragraphs import _apply_paragraph_style
            from .html_docx_blocks import _ensure_paragraph

            # Cell borders are applied at the cell level (w:tcBorders). Strip
            # border-* from the style passed to paragraph content so it doesn't
            # also produce paragraph-level borders (w:pBdr), which would render
            # as a second visible line on top of the cell border.
            content_style = {k: v for k, v in cell_style.items() if not k.startswith("border")}

            if _has_block_children(cell_node):
                for child in cell_node.children:
                    _handle_block(child, docx_cell, css_rules, content_style, table_auto_widths)
            else:
                # Cell contains only inline content (text, spans, <br>) —
                # render as one paragraph so <br> becomes a soft line break
                # rather than each text node becoming a separate paragraph.
                paragraph = _ensure_paragraph(docx_cell)
                _apply_paragraph_style(paragraph, content_style)
                _add_inline_runs(cell_node, paragraph, css_rules, content_style)

            col_idx += max(colspan, 1)

    if is_collapsed:
        inside_h = collapsed_borders.get("inside_h")
        inside_v = collapsed_borders.get("inside_v")
        if inside_h or inside_v:
            _apply_collapsed_table_borders(
                docx_table,
                {"inside_h": inside_h, "inside_v": inside_v},
            )


def _extract_table_column_widths(
    table_node: Tag,
    table_style: Dict[str, str],
    column_count: int,
    table_auto_widths=None,
) -> Optional[List[float]]:
    if column_count <= 0:
        return None
    rows = table_node.find_all("tr", recursive=False)
    if not rows:
        for section in table_node.find_all(["thead", "tbody", "tfoot"], recursive=False):
            rows.extend(section.find_all("tr", recursive=False))
    if not rows:
        return None

    widths: List[Optional[float]] | None = None
    for row in rows:
        cells = [cell for cell in row.find_all(["td", "th"], recursive=False)]
        if not cells:
            continue
        candidate_widths: List[Optional[float]] = []
        for cell in cells:
            style_map = _compute_style_map(cell, [])
            width_value = style_map.get("width", "") or cell.get("width", "")
            if isinstance(width_value, str) and width_value.endswith("%"):
                candidate_widths.append(_parse_percentage(width_value))
            else:
                candidate_widths.append(_parse_px_to_pt(str(width_value)) if width_value else None)
        if any(candidate_widths) and len(candidate_widths) >= column_count:
            widths = candidate_widths
            break

    if not widths or not any(widths):
        return None

    total_width_pt = _parse_px_to_pt(table_style.get("width", ""))
    if total_width_pt is None:
        # width:100% or a CSS percentage — look up the Playwright-rendered width.
        # Without the actual rendered width, percentage column values can't be
        # resolved correctly (the fallback of 547pt is wrong for nested tables).
        auto_id = table_node.get("data-table-id")
        if auto_id and table_auto_widths and auto_id in table_auto_widths:
            auto = table_auto_widths[auto_id]
            if auto:
                total_width_pt = sum(auto)
        if total_width_pt is None:
            # No context available — let _extract_auto_widths handle this table.
            return None

    resolved: List[float | None] = []
    fixed_total = sum(w for w in widths if isinstance(w, float) and w > 1.0)
    remaining = max(total_width_pt - fixed_total, 0)

    for width in widths:
        if width is None:
            resolved.append(None)
        elif width <= 1.0:
            # Treat as an absolute fraction of total_width_pt, not normalised
            # against other percentage columns. This correctly handles mixed
            # rows where some columns specify a % and others have no width.
            resolved.append(total_width_pt * width)
        else:
            resolved.append(width)

    none_count = sum(1 for value in resolved if value is None)
    if none_count:
        filled_total = sum(value for value in resolved if isinstance(value, float))
        filler = max(total_width_pt - filled_total, 0) / none_count
        resolved = [value if value is not None else filler for value in resolved]

    if len(resolved) < column_count:
        missing = column_count - len(resolved)
        filler = (total_width_pt - sum(resolved)) / max(missing, 1)
        resolved.extend([filler] * missing)
    return [float(value) for value in resolved[:column_count]]


def _apply_table_column_widths(table, widths_pt: List[float]) -> None:
    widths_pt = _adjust_column_widths_for_parent_padding(table, widths_pt)
    widths_pt = _adjust_column_widths_for_outer_borders(table, widths_pt)
    _update_table_grid(table, widths_pt)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_pt) and widths_pt[idx] > 0:
                _set_cell_width(cell, widths_pt[idx])


def _update_table_grid(table, widths_pt: List[float]) -> None:
    """Replace w:tblGrid with column definitions matching widths_pt.

    python-docx creates w:tblGrid with equal-width columns when add_table() is
    called. Without updating it, Word ignores the per-cell w:tcW values and
    renders equal columns. This must be called before setting cell widths.
    """
    tbl = table._tbl
    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_pt:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 20)))
        grid.append(col)
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is not None:
        tbl_pr.addnext(grid)
    else:
        tbl.insert(0, grid)


def _set_cell_width(cell, width_pt: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_pt * 20)))


def _set_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tc_pr.append(no_wrap)


def _estimate_right_column_width_pt(column_widths_pt: Optional[List[float]]) -> float:
    if column_widths_pt and len(column_widths_pt) >= 2:
        return max(column_widths_pt[1], 160.0)
    return 180.0


def _merge_styles(parent_style: Dict[str, str], own_style: Dict[str, str]) -> Dict[str, str]:
    merged: Dict[str, str] = {}
    for key in _INHERITABLE_STYLES:
        if key in parent_style:
            merged[key] = parent_style[key]
    for key, value in own_style.items():
        merged[key] = value
    return merged


def _apply_table_styles(table, style_map: Dict[str, str]) -> None:
    width_value = style_map.get("width", "") or style_map.get("max-width", "")
    width_pt = _parse_px_to_pt(width_value)
    if width_pt:
        _set_table_width(table, width_pt)

    if _should_center_table(style_map):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _set_table_normal_style(table)
    _clear_table_look(table)
    _set_table_default_cell_margins(table, 0, 0, 0, 0)
    _set_table_cell_spacing(table, 0)
    _apply_table_border(table, style_map)


def _set_table_width(table, width_pt: float) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.append(tbl_pr)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_pt * 20)))


def _apply_table_border(table, style_map: Dict[str, str]) -> None:
    border = style_map.get("border", "")
    if not border:
        return
    parts = border.split()
    width_pt = _parse_px_to_pt(parts[0]) if parts else None
    color = None
    for part in parts:
        if part.startswith("#"):
            color = part[1:].upper()
            break
        if part.lower() in _NAMED_COLORS:
            r, g, b = _NAMED_COLORS[part.lower()]
            color = f"{r:02X}{g:02X}{b:02X}"
            break
    if not width_pt or not color:
        return
    _apply_table_outer_borders(table, width_pt, color)
    table._docs_border_width_pt = width_pt


def _set_table_normal_style(table) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.append(tbl_pr)
    tbl_style = tbl_pr.find(qn("w:tblStyle"))
    if tbl_style is None:
        tbl_style = OxmlElement("w:tblStyle")
        tbl_pr.insert(0, tbl_style)
    tbl_style.set(qn("w:val"), "TableNormal")


def _clear_table_look(table) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return
    tbl_look = tbl_pr.find(qn("w:tblLook"))
    if tbl_look is not None:
        tbl_pr.remove(tbl_look)


def _set_table_default_cell_margins(
    table, top: float, right: float, bottom: float, left: float
) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.append(tbl_pr)
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for side, value in [("top", top), ("right", right), ("bottom", bottom), ("left", left)]:
        elem = tbl_cell_mar.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(elem)
        elem.set(qn("w:w"), str(int(value * 20)))
        elem.set(qn("w:type"), "dxa")


def _set_table_cell_spacing(table, spacing_pt: float) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.append(tbl_pr)
    tbl_cell_spacing = tbl_pr.find(qn("w:tblCellSpacing"))
    if spacing_pt <= 0:
        if tbl_cell_spacing is not None:
            tbl_pr.remove(tbl_cell_spacing)
        return
    if tbl_cell_spacing is None:
        tbl_cell_spacing = OxmlElement("w:tblCellSpacing")
        tbl_pr.append(tbl_cell_spacing)
    tbl_cell_spacing.set(qn("w:w"), str(int(spacing_pt * 20)))
    tbl_cell_spacing.set(qn("w:type"), "dxa")


def _should_center_table(style_map: Dict[str, str]) -> bool:
    margin = style_map.get("margin", "")
    if margin and "auto" in margin:
        return True
    left = style_map.get("margin-left", "").strip().lower()
    right = style_map.get("margin-right", "").strip().lower()
    if left == "auto" and right == "auto":
        return True
    return False


def _apply_cell_styles(
    cell,
    style_map: Dict[str, str],
    apply_padding: bool = True,
    suppress_borders: Optional[set] = None,
) -> None:
    bg_color = _parse_background_color(style_map)
    if bg_color:
        _set_cell_shading(cell, bg_color)

    _apply_cell_vertical_alignment(cell, style_map)

    suppress = suppress_borders or set()
    border = _parse_border(style_map.get("border", ""))
    top_border = _parse_border(style_map.get("border-top", ""))
    right_border = _parse_border(style_map.get("border-right", ""))
    bottom_border = _parse_border(style_map.get("border-bottom", ""))
    left_border = _parse_border(style_map.get("border-left", ""))

    effective_top    = (top_border    or border) if "top"    not in suppress else None
    effective_right  = (right_border  or border) if "right"  not in suppress else None
    effective_bottom = (bottom_border or border) if "bottom" not in suppress else None
    effective_left   = (left_border   or border) if "left"   not in suppress else None

    if any([effective_top, effective_right, effective_bottom, effective_left]):
        _set_cell_border(
            cell,
            top=effective_top,
            right=effective_right,
            bottom=effective_bottom,
            left=effective_left,
        )

    if "left" not in suppress:
        border_left = _parse_border_left(style_map)
        if border_left:
            width_pt, color = border_left
            _set_cell_border(cell, left=(width_pt, color))

    padding = _resolve_padding(style_map)
    if padding and apply_padding:
        top, right, bottom, left = _normalize_padding(padding)
        _set_cell_margins(
            cell,
            top * _PADDING_SCALE,
            right * _PADDING_SCALE,
            bottom * _PADDING_SCALE,
            left * _PADDING_SCALE,
        )
    else:
        _set_cell_margins(cell, 0, 0, 0, 0)


def _apply_cell_vertical_alignment(cell, style_map: Dict[str, str]) -> None:
    vertical_align = style_map.get("vertical-align", "").strip().lower()
    if not vertical_align:
        return
    if vertical_align in {"middle", "center"}:
        val = "center"
    elif vertical_align in {"top", "bottom"}:
        val = vertical_align
    else:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is None:
        v_align = OxmlElement("w:vAlign")
        tc_pr.append(v_align)
    v_align.set(qn("w:val"), val)


def _apply_cell_padding_spacing(cell, padding: Tuple[float, float, float, float]) -> None:
    paragraphs = [p for p in cell.paragraphs if _is_direct_cell_paragraph(cell, p)]
    if not paragraphs:
        return
    top, right, bottom, left = _normalize_padding(padding)
    if top:
        _add_paragraph_spacing(paragraphs[0], before_pt=top * _PADDING_SCALE)
    if bottom:
        _add_paragraph_spacing(paragraphs[-1], after_pt=bottom * _PADDING_SCALE)
    if left or right:
            for paragraph in paragraphs:
                if paragraph.style and paragraph.style.name == "List Bullet":
                    # "List Bullet" indentation is controlled by the numbering definition;
                    # adding paragraph-level w:ind here overrides it and misaligns bullets.
                    pass
                else:
                    _add_paragraph_indent(
                        paragraph,
                        left_pt=left * _PADDING_SCALE,
                        right_pt=right * _PADDING_SCALE,
                    )


def _cell_has_only_tables(cell) -> bool:
    try:
        if not cell.tables:
            return False
        for paragraph in cell.paragraphs:
            if paragraph.text.strip():
                return False
        return True
    except Exception:
        return "<w:tbl>" in cell._tc.xml


def _is_direct_cell_paragraph(cell, paragraph) -> bool:
    try:
        return paragraph._p.getparent() is cell._tc
    except Exception:
        return True


def _apply_nested_table_vertical_padding(
    cell, padding: Tuple[float, float, float, float]
) -> None:
    if not cell.tables:
        return
    top, _right, bottom, _left = _normalize_padding(padding)
    top_pt = top * _PADDING_SCALE if top else 0
    bottom_pt = bottom * _PADDING_SCALE if bottom else 0
    tables = cell.tables
    if top_pt and tables:
        first_table = tables[0]
        if first_table.rows:
            for tcell in first_table.rows[0].cells:
                _add_cell_margins(tcell, top=top_pt)
    if bottom_pt and tables:
        last_table = tables[-1]
        if last_table.rows:
            last_row = last_table.rows[-1]
            for tcell in last_row.cells:
                _add_cell_margins(tcell, bottom=bottom_pt)


def _set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_border(
    cell,
    top: Optional[Tuple[float, str]] = None,
    right: Optional[Tuple[float, str]] = None,
    bottom: Optional[Tuple[float, str]] = None,
    left: Optional[Tuple[float, str]] = None,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for side_name, val in [
        ("w:top", top),
        ("w:right", right),
        ("w:bottom", bottom),
        ("w:left", left),
    ]:
        if not val:
            continue
        width_pt, color = val
        # Remove any existing element for this side to prevent duplicates when
        # _set_cell_border is called more than once on the same cell.
        for existing in borders.findall(qn(side_name)):
            borders.remove(existing)
        el = OxmlElement(side_name)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), _border_sz(width_pt))
        el.set(qn("w:color"), color)
        borders.append(el)


def _apply_table_outer_borders(table, width_pt: float, color: str) -> None:
    if not table.rows:
        return
    last_row_idx = len(table.rows) - 1
    last_col_idx = len(table.rows[0].cells) - 1
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            top = (width_pt, color) if row_idx == 0 else None
            bottom = (width_pt, color) if row_idx == last_row_idx else None
            left = (width_pt, color) if col_idx == 0 else None
            right = (width_pt, color) if col_idx == last_col_idx else None
            if top or right or bottom or left:
                _set_cell_border(
                    cell,
                    top=top,
                    right=right,
                    bottom=bottom,
                    left=left,
                )


def _adjust_column_widths_for_outer_borders(
    table, widths_pt: List[float]
) -> List[float]:
    border_width = getattr(table, "_docs_border_width_pt", 0)
    if border_width <= 0 or not widths_pt:
        return widths_pt
    adjusted = list(widths_pt)
    if len(adjusted) == 1:
        adjusted[0] = max(adjusted[0] - (border_width * 2), 0)
        return adjusted
    adjusted[0] = max(adjusted[0] - border_width, 0)
    adjusted[-1] = max(adjusted[-1] - border_width, 0)
    return adjusted


def _apply_table_parent_padding(table, parent_style: Dict[str, str]) -> None:
    padding = _resolve_padding(parent_style)
    if not padding:
        return
    top, right, bottom, left = _normalize_padding(padding)
    left_pt = left * _PADDING_SCALE
    right_pt = right * _PADDING_SCALE
    if left_pt:
        _set_table_indent(table, left_pt)
    table._docs_parent_padding_pt = (
        left_pt,
        right_pt,
        top * _PADDING_SCALE,
        bottom * _PADDING_SCALE,
    )


def _set_table_indent(table, indent_pt: float) -> None:
    if indent_pt <= 0:
        return
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.append(tbl_pr)
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(int(indent_pt * 20)))
    tbl_ind.set(qn("w:type"), "dxa")


def _adjust_column_widths_for_parent_padding(
    table, widths_pt: List[float]
) -> List[float]:
    padding = getattr(table, "_docs_parent_padding_pt", None)
    if not padding:
        return widths_pt
    left_pt, right_pt, _top_pt, _bottom_pt = padding
    shrink = max(left_pt + right_pt, 0)
    if shrink <= 0:
        return widths_pt
    total = sum(widths_pt)
    if total <= 0 or total <= shrink:
        return widths_pt
    scale = (total - shrink) / total
    return [max(width * scale, 0) for width in widths_pt]


def _set_cell_margins(cell, top: float, right: float, bottom: float, left: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for side, value in [("top", top), ("right", right), ("bottom", bottom), ("left", left)]:
        elem = tc_mar.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            tc_mar.append(elem)
        elem.set(qn("w:w"), str(int(value * 20)))
        elem.set(qn("w:type"), "dxa")


def _should_prevent_row_split(table_node: Tag, style_map: Dict[str, str]) -> bool:
    """Return True if w:cantSplit should be applied to this table's rows.

    Triggers when:
    - auto_page_breaks marked the table as fitting within one page
      (data-docx-cant-split attribute), OR
    - the agent explicitly set page-break-inside:avoid on the table.
    """
    if table_node.get("data-docx-cant-split"):
        return True
    for key in ("page-break-inside", "break-inside"):
        if style_map.get(key, "").strip().lower() == "avoid":
            return True
    return False


def _set_table_cant_split(table) -> None:
    """Add w:cantSplit to every row so Word won't split a row across pages."""
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)


def _collect_collapsed_borders(
    row_cells: list[list],
    row_colspans: list[list[int]],
    max_cols: int,
    css_rules,
) -> dict[str, Optional[tuple[float, str]]]:
    """Collect all six border positions for a border-collapse:collapse table.

    Returns a dict with keys: inside_h, inside_v, top, bottom, left, right.
    Each value is (width_pt, color_hex) or None.

    The complete set is needed so _apply_collapsed_table_borders can write all six
    directions into w:tblBorders without mixing in any cell-level tcBorders.
    """

    def first_cell_border(cells, *props) -> Optional[tuple[float, str]]:
        for cell_node in cells:
            style = _compute_style_map(cell_node, css_rules)
            for prop in props:
                val = _parse_border(style.get(prop, ""))
                if val:
                    return val
            # Fall back to shorthand border
            val = _parse_border(style.get("border", ""))
            if val:
                return val
        return None

    inside_h: Optional[tuple[float, str]] = None
    inside_v: Optional[tuple[float, str]] = None

    # interior-H: border-bottom on non-last rows, then border-top on non-first rows
    for cells in row_cells[:-1]:
        inside_h = first_cell_border(cells, "border-bottom")
        if inside_h:
            break
    if not inside_h and len(row_cells) > 1:
        for cells in row_cells[1:]:
            inside_h = first_cell_border(cells, "border-top")
            if inside_h:
                break

    # interior-V: border-right on non-last cols, then border-left on non-first cols
    if max_cols > 1:
        for row_idx, cells in enumerate(row_cells):
            colspans = row_colspans[row_idx]
            col_pos = 0
            for cell_idx, cell_node in enumerate(cells):
                span = colspans[cell_idx] if cell_idx < len(colspans) else 1
                if col_pos + span < max_cols:
                    style = _compute_style_map(cell_node, css_rules)
                    inside_v = _parse_border(style.get("border-right", "")) or _parse_border(
                        style.get("border", "")
                    )
                    if inside_v:
                        break
                col_pos += span
            if inside_v:
                break

        if not inside_v:
            for row_idx, cells in enumerate(row_cells):
                colspans = row_colspans[row_idx]
                col_pos = 0
                for cell_idx, cell_node in enumerate(cells):
                    span = colspans[cell_idx] if cell_idx < len(colspans) else 1
                    if col_pos > 0:
                        style = _compute_style_map(cell_node, css_rules)
                        inside_v = _parse_border(style.get("border-left", ""))
                        if inside_v:
                            break
                    col_pos += span
                if inside_v:
                    break

    # outer edges — collect from the boundary cells
    outer_top = first_cell_border(row_cells[0], "border-top") if row_cells else None
    outer_bottom = first_cell_border(row_cells[-1], "border-bottom") if row_cells else None

    first_col_cells = [cells[0] for cells in row_cells if cells]
    outer_left = first_cell_border(first_col_cells, "border-left") if first_col_cells else None

    def last_col_cell(cells, colspans):
        col_pos = 0
        last = None
        for idx, cell_node in enumerate(cells):
            span = colspans[idx] if idx < len(colspans) else 1
            last = cell_node
            col_pos += span
        return last

    last_col_cells = [
        last_col_cell(cells, row_colspans[ri]) for ri, cells in enumerate(row_cells) if cells
    ]
    last_col_cells = [c for c in last_col_cells if c is not None]
    outer_right = first_cell_border(last_col_cells, "border-right") if last_col_cells else None

    return {
        "inside_h": inside_h,
        "inside_v": inside_v,
        "top": outer_top,
        "bottom": outer_bottom,
        "left": outer_left,
        "right": outer_right,
    }


def _apply_collapsed_table_borders(
    table,
    collapsed_borders: dict[str, Optional[tuple[float, str]]],
) -> None:
    """Write all six border positions into w:tblBorders for a border-collapse:collapse table.

    Using tblBorders exclusively (no tcBorders on any cell) guarantees that Word
    draws exactly one line per edge and never mixes table-level and cell-level borders,
    which is the root cause of double-line rendering.
    """
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.append(tbl_pr)
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    # OOXML requires a specific child order inside w:tblBorders.
    ordered_sides = [
        ("w:top", collapsed_borders.get("top")),
        ("w:left", collapsed_borders.get("left")),
        ("w:bottom", collapsed_borders.get("bottom")),
        ("w:right", collapsed_borders.get("right")),
        ("w:insideH", collapsed_borders.get("inside_h")),
        ("w:insideV", collapsed_borders.get("inside_v")),
    ]
    # Clear existing children to guarantee order and avoid stale entries.
    for child in list(borders):
        borders.remove(child)

    for tag_name, border_val in ordered_sides:
        if not border_val:
            continue
        width_pt, color = border_val
        el = OxmlElement(tag_name)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), _border_sz(width_pt))
        el.set(qn("w:color"), color)
        borders.append(el)


def _add_cell_margins(
    cell,
    top: Optional[float] = None,
    right: Optional[float] = None,
    bottom: Optional[float] = None,
    left: Optional[float] = None,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for side, value in [("top", top), ("right", right), ("bottom", bottom), ("left", left)]:
        if value is None:
            continue
        elem = tc_mar.find(qn(f"w:{side}"))
        current_val = 0
        if elem is not None:
            current = elem.get(qn("w:w"))
            if current and current.isdigit():
                current_val = int(current)
        else:
            elem = OxmlElement(f"w:{side}")
            tc_mar.append(elem)
        elem.set(qn("w:w"), str(current_val + int(value * 20)))
        elem.set(qn("w:type"), "dxa")
