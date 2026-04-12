import base64
import binascii
import mimetypes
import re
from io import BytesIO
from pathlib import Path
from typing import Dict, Optional
from urllib.request import Request, urlopen

from bs4.element import Tag
from cairosvg import svg2png
from docx.shared import Pt

from .html_docx_css import _parse_length_to_pt
from .html_docx_selectors import _compute_style_map

_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg", ".ico", ".bmp", ".avif"}

# Fallback max image width when no cell/container constraint is available.
# ~5.5 inches — fits safely within any standard A4/Letter body.
_DEFAULT_MAX_IMAGE_WIDTH_PT = 400.0


def _is_image_path(src: str) -> bool:
    return Path(src.split("?")[0]).suffix.lower() in _IMAGE_EXTENSIONS


def embed_local_images(html: str, base_dir: Path) -> str:
    """Rewrite all image references to inline base64 data URIs.

    Handles:
    - HTML src= attributes (<img>, <image>)
    - CSS url() background images
    - SVG href= / xlink:href= on <image> elements
    - <object data=> embeds

    Local paths are resolved relative to base_dir.
    HTTP/HTTPS URLs are fetched and embedded so documents work offline.
    data: URIs, file:// URLs, fragment-only (#) refs, non-image paths, and
    unreachable resources are left unchanged (no crash).
    """

    def _encode(src: str) -> str | None:
        if src.startswith("data:") or src.startswith("file://"):
            return None

        if src.startswith("http://") or src.startswith("https://"):
            try:
                req = Request(src, headers={"User-Agent": "docs-agent"})
                with urlopen(req, timeout=20) as response:
                    data = response.read()
                mime, _ = mimetypes.guess_type(src.split("?")[0])
                mime = mime or "image/png"
                return f"data:{mime};base64,{base64.b64encode(data).decode('ascii')}"
            except Exception:
                return None

        if not _is_image_path(src):
            return None

        img_path = (base_dir / src).resolve()
        if not img_path.exists():
            return None
        mime, _ = mimetypes.guess_type(str(img_path))
        mime = mime or "image/png"
        return f"data:{mime};base64,{base64.b64encode(img_path.read_bytes()).decode('ascii')}"

    def replace_src(match: re.Match) -> str:
        quote, src = match.group(1), match.group(2)
        data_uri = _encode(src)
        return f"src={quote}{data_uri}{quote}" if data_uri else match.group(0)

    def replace_css_url(match: re.Match) -> str:
        quote, src = match.group(1), match.group(2)
        data_uri = _encode(src)
        return f"url({quote}{data_uri}{quote})" if data_uri else match.group(0)

    def replace_href(match: re.Match) -> str:
        attr, quote, src = match.group(1), match.group(2), match.group(3)
        data_uri = _encode(src)
        return f"{attr}={quote}{data_uri}{quote}" if data_uri else match.group(0)

    html = re.sub(r'src=(["\'])((?!data:|https?://|file://)[^"\']+)\1', replace_src, html)
    html = re.sub(r'src=(["\'])(https?://[^"\']+)\1', replace_src, html)
    html = re.sub(r"url\(([\"']?)((?!data:|file://)[^\"')>\s]+)\1\)", replace_css_url, html)
    html = re.sub(
        r'(href|xlink:href|data)=(["\'])((?!data:|https?://|file://|#)[^"\']+)\2',
        replace_href,
        html,
    )
    return html


def _add_svg_run(paragraph, node: Tag, parent_style: Dict[str, str]) -> None:
    """Rasterize an inline <svg> node to PNG and add it as a picture run.

    BeautifulSoup's html.parser has two SVG serialization bugs:
    - Lowercases attribute names (viewBox → viewbox), breaking cairosvg.
    - Serializes SVG-specific attrs like 'fill' as valueless booleans, making
      paths transparent.

    We avoid both by reconstructing the SVG XML from the parsed .attrs dicts
    directly rather than using str(node). No HTML string manipulation.
    """
    style = node.get("style", "") or ""
    m = re.search(r"width\s*:\s*([\d.]+\s*(?:pt|px|em|rem)?)", style, re.IGNORECASE)
    width_pt = _parse_length_to_pt(m.group(1)) if m else None
    output_width = max(100, int((width_pt or 54) * (96 / 72) * 2))  # 2× for retina

    svg_xml = _svg_node_to_xml(node, output_width)

    try:
        png_bytes = svg2png(bytestring=svg_xml.encode("utf-8"))
    except Exception:
        return  # silently skip if conversion fails

    run = paragraph.add_run()
    _add_picture_safe(run, BytesIO(png_bytes), width_pt, None)


def _svg_node_to_xml(node: Tag, output_width: int) -> str:
    """Reconstruct SVG XML from a BeautifulSoup node, bypassing html.parser serialization bugs.

    html.parser drops 'fill' values and lowercases 'viewBox'. We build the XML
    directly from node.attrs (which are always correct) and recurse into children.
    Explicit width/height are injected on the root <svg> so cairosvg can render.
    """
    vb = node.get("viewbox", "")  # BeautifulSoup lowercases to 'viewbox'
    parts = re.split(r"[\s,]+", vb.strip())
    if len(parts) == 4:
        try:
            vb_w, vb_h = float(parts[2]), float(parts[3])
            aspect = vb_h / vb_w if vb_w else 1.0
        except ValueError:
            aspect = 1.0
    else:
        aspect = 1.0
    output_height = max(1, int(output_width * aspect))

    ns = node.get("xmlns", "http://www.w3.org/2000/svg")
    children_xml = "".join(_svg_child_to_xml(c) for c in node.children)
    return (
        f'<svg xmlns="{ns}" viewBox="{vb}" '
        f'width="{output_width}" height="{output_height}">'
        f"{children_xml}</svg>"
    )


def _svg_child_to_xml(node) -> str:
    """Recursively serialize an SVG child node to XML, preserving all attribute values."""
    from bs4.element import NavigableString
    if isinstance(node, NavigableString):
        return str(node)
    if not isinstance(node, Tag):
        return ""
    attrs = " ".join(f'{k}="{v}"' for k, v in node.attrs.items())
    inner = "".join(_svg_child_to_xml(c) for c in node.children)
    if inner:
        return f"<{node.name} {attrs}>{inner}</{node.name}>" if attrs else f"<{node.name}>{inner}</{node.name}>"
    return f"<{node.name} {attrs}/>" if attrs else f"<{node.name}/>"


def _add_image_run(paragraph, node: Tag, parent_style: Dict[str, str]) -> None:
    src = node.get("src", "") or ""
    if not src:
        return

    width_pt, height_pt = _extract_image_dimensions(node, parent_style)
    image_stream = _load_image_stream(src)
    if image_stream is None:
        alt_text = (node.get("alt", "") or "").strip()
        if alt_text:
            paragraph.add_run(alt_text)
        return

    run = paragraph.add_run()
    if width_pt and height_pt:
        _add_picture_safe(run, image_stream, width_pt, height_pt)
    elif width_pt:
        _add_picture_safe(run, image_stream, width_pt, None)
    elif height_pt:
        _add_picture_safe(run, image_stream, None, height_pt)
    else:
        _add_picture_safe(run, image_stream, None, None)


def _extract_image_dimensions(
    node: Tag, parent_style: Dict[str, str]
) -> tuple[Optional[float], Optional[float]]:
    style_map = _compute_style_map(node, [])
    width_value = style_map.get("width", "") or node.get("width", "")
    height_value = style_map.get("height", "") or node.get("height", "")

    width_pt = _parse_length_to_pt(str(width_value)) if width_value else None
    height_pt = _parse_length_to_pt(str(height_value)) if height_value else None

    if width_pt is None:
        width_pt = (
            _parse_length_to_pt(str(parent_style.get("width", "")))
            if parent_style.get("width")
            else None
        )
    if height_pt is None:
        height_pt = (
            _parse_length_to_pt(str(parent_style.get("height", "")))
            if parent_style.get("height")
            else None
        )

    # Determine the available width cap: prefer the cell width injected by the
    # table processor; fall back to a page-body constant for images outside tables.
    cell_width_str = parent_style.get("_cell_width_pt", "")
    max_width_pt = float(cell_width_str) if cell_width_str else _DEFAULT_MAX_IMAGE_WIDTH_PT

    if width_pt is None:
        # No explicit width (e.g. width:100%) — fill available space and let
        # aspect ratio determine height.
        width_pt = max_width_pt
        height_pt = None
    elif width_pt > max_width_pt:
        # Explicit width exceeds the available space — cap and preserve aspect ratio.
        width_pt = max_width_pt
        height_pt = None

    return width_pt, height_pt


def _load_image_stream(src: str) -> BytesIO | Path | None:
    if src.startswith("data:image/"):
        return _load_data_image(src)

    if src.startswith("http://") or src.startswith("https://"):
        try:
            request = Request(src, headers={"User-Agent": "docs-agent"})
            with urlopen(request, timeout=20) as response:
                return BytesIO(response.read())
        except Exception:
            return None

    path = Path(src)
    if path.exists():
        if path.suffix.lower() == ".svg":
            return _convert_svg_to_png(path.read_bytes())
        return path
    return None


def _load_data_image(src: str) -> BytesIO | None:
    try:
        header, encoded = src.split(",", 1)
    except ValueError:
        return None
    is_base64 = "base64" in header
    if "image/svg+xml" in header:
        svg_bytes = _decode_data_uri(encoded, is_base64)
        return _convert_svg_to_png(svg_bytes)
    data = _decode_data_uri(encoded, is_base64)
    return BytesIO(data) if data else None


def _decode_data_uri(encoded: str, is_base64: bool) -> Optional[bytes]:
    try:
        if is_base64:
            return base64.b64decode(encoded)
        return encoded.encode("utf-8")
    except (binascii.Error, UnicodeEncodeError):
        return None


def _convert_svg_to_png(svg_bytes: Optional[bytes]) -> BytesIO | None:
    if not svg_bytes:
        return None
    try:
        png_bytes = svg2png(bytestring=svg_bytes)
        return BytesIO(png_bytes)
    except Exception:
        return None


def _add_picture_safe(
    run, image_stream, width_pt: Optional[float], height_pt: Optional[float]
) -> None:
    try:
        if width_pt is not None and height_pt is not None:
            run.add_picture(image_stream, width=Pt(width_pt), height=Pt(height_pt))
        elif width_pt is not None:
            run.add_picture(image_stream, width=Pt(width_pt))
        elif height_pt is not None:
            run.add_picture(image_stream, height=Pt(height_pt))
        else:
            run.add_picture(image_stream)
    except Exception:
        run.add_text("[image]")
